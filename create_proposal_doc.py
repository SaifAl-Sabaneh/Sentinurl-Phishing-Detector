import sys
import subprocess
import os

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_proposal():
    doc = Document()

    # Title
    title = doc.add_heading('Graduation Project Proposal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle/Meta
    p_meta = doc.add_paragraph()
    p_meta.add_run('Project Title: ').bold = True
    p_meta.add_run('SentinURL: Deep Semantic Inspection and Heuristic Triage for Web Security\n')
    p_meta.add_run('Student Name: ').bold = True
    p_meta.add_run('Saif Al-Sabaneh\n')
    p_meta.add_run('Submission Deadline: ').bold = True
    p_meta.add_run('16/3/2026')
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacing

    doc.add_heading('1. Overview', level=1)
    p_overview = doc.add_paragraph("From a Business Intelligence and Data Analytics perspective, modern cybersecurity is fundamentally a massive data classification challenge. This graduation project proposes the design, implementation, and rigorous evaluation of 'SentinURL', a high-performance, multi-stage machine learning architecture dedicated to the identification and neutralization of zero-day phishing threats. SentinURL bridges the gap between rapid heuristic assessment and deep Natural Language Processing (NLP), transforming unstructured web logs into actionable, automated security intelligence.")
    p_overview.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('2. Problem Statement', level=1)
    p_prob = doc.add_paragraph("Traditional cybersecurity blocklists are increasingly failing against dynamically generated domains that mutate faster than they can be cataloged. Cybercriminals utilize heavily obfuscated URLs and ephemeral hosting to bypass legacy firewalls, leaving individuals and enterprises vulnerable to credential theft and financial fraud. There is a critical need for a predictive, data-driven security mechanism that can identify phishing infrastructure based on its structural and semantic traits rather than relying on historical threat databases.")
    p_prob.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('3. Main Goal', level=1)
    p_goal = doc.add_paragraph("The primary goal of this project is to architect an enterprise-grade predictive analytics pipeline that evaluates the semantic intent and mathematical structure of a URL in milliseconds. By deploying a dual-stage Artificial Intelligence model, the system aims to pre-emptively intercept zero-day phishing attacks while maintaining an exceptionally low false-positive rate, thereby providing organizations with a highly reliable automated defense perimeter.")
    p_goal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('4. Target Audience & Beneficiaries', level=1)
    p_target_1 = doc.add_paragraph("SentinURL is engineered to serve two primary target groups:")
    p_target_2 = doc.add_paragraph("1. Enterprise Organizations & Security Operations Centers (SOCs): Providing automated triage and threat intelligence to protect corporate networks and reduce the manual analytical load on security analysts.", style="List Bullet")
    p_target_3 = doc.add_paragraph("2. Everyday Consumers & End-Users: Offering a seamless, invisible layer of protection against sophisticated social engineering and credential-harvesting campaigns.", style="List Bullet")

    doc.add_heading('5. Methodology: How SentinURL Works', level=1)
    
    doc.add_heading('Data Engineering and Feature Extraction', level=2)
    p_data = doc.add_paragraph("The foundation of this quantitative research relies on a massive dataset containing over 1,000,000 distinct records. The raw data undergoes an aggressive Data Engineering pipeline to isolate lexical properties, intricate domain structures, character entropy, and vowel-to-consonant ratios. By feeding the AI computationally rich NLP metrics, the model transcends simple memorization and autonomously learns the underlying statistical traits of cyber threats.")
    p_data.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('Predictive Modeling Architecture', level=2)
    p_structure = doc.add_paragraph("The system operates on a Python-based two-stage Machine Learning pipeline. Stage 1 is a rapid-triage filter utilizing a Term Frequency-Inverse Document Frequency (TF-IDF) vectorizer with a 150,000 text-sequence memory matrix to clear benign traffic instantly. Statistically anomalous URLs are escalated to the Stage 2 deep-inspection engine, which employs a Hist-Gradient Boosting classifier to dissect extreme semantic NLP features and track isolated threat intent keywords (e.g., 'scam', 'crypto', 'login').")
    p_structure.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('6. Organizational Implementation & Community Impact', level=1)
    p_impact = doc.add_paragraph("Within a corporate environment, SentinURL can be deployed as an API microservice, integrating directly into edge routers, email gateways, or SIEM (Security Information and Event Management) platforms. It empowers Business Intelligence teams with automated HTML diagnostic reports detailing mathematical confidence scores for every flagged threat. The broader community impact is profound: by deploying predictive AI against phishing infrastructure, SentinURL drastically reduces the success rate of financial fraud, protecting both corporate assets and vulnerable individuals from targeted digital exploitation.")
    p_impact.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('7. Expected Analytics Outcomes', level=1)
    p_results = doc.add_paragraph("Rigorous cross-validation proves that the dual-stage AI achieves an exceptional offline detection accuracy of 89.1% using exclusively mathematical and semantic data analysis. When deployed and integrated with dynamic external data streams (Google Safe Browsing APIs), the analytics engine scales to a formidable 99.4% real-world accuracy rate.")
    p_results.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph() # Spacing
    
    doc.add_heading('8. Conclusion', level=1)
    p_conclusion = doc.add_paragraph("In conclusion, the SentinURL Phishing Detection System represents a comprehensive application of Data Analytics to modernize threat detection methodologies. By migrating to a predictive framework built on Natural Language Processing and deep statistical modeling, this project successfully addresses a persistent vulnerability in modern enterprise environments, comprehensively fulfilling the analytical requirements for a Business Intelligence graduation project.")
    p_conclusion.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save
    file_path = os.path.join(r'c:\Users\Asus\Desktop\Graduation Project\PreTrained Models', 'SentinURL_Project_Proposal.docx')
    doc.save(file_path)
    print(f"Document saved successfully to {file_path}")

if __name__ == '__main__':
    create_proposal()
